import os
from typing import Optional
import PyPDF2
from docx import Document


class ResumeParser:
    """Parse text content from PDF and DOCX resume files."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Optional[str]:
        """Extract text from a PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading PDF file {file_path}: {e}")
            return None

    @staticmethod
    def extract_text_from_pdf_bytes(file_content: bytes) -> Optional[str]:
        """Extract text from PDF bytes in memory."""
        try:
            import io
            text = ""
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading PDF from bytes: {e}")
            return None

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {e}")
            return None

    @staticmethod
    def extract_text_from_docx_bytes(file_content: bytes) -> Optional[str]:
        """Extract text from DOCX bytes in memory."""
        try:
            import io
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX from bytes: {e}")
            return None

    @classmethod
    def parse_resume(cls, file_path: str) -> Optional[str]:
        """Parse resume file and extract text content."""
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return cls.extract_text_from_docx(file_path)
        else:
            print(f"Unsupported file format: {file_extension}")
            return None

    @classmethod
    def parse_resume_from_bytes(cls, file_content: bytes, filename: str) -> Optional[str]:
        """Parse resume from bytes in memory."""
        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension == '.pdf':
            return cls.extract_text_from_pdf_bytes(file_content)
        elif file_extension in ['.docx', '.doc']:
            return cls.extract_text_from_docx_bytes(file_content)
        else:
            print(f"Unsupported file format: {file_extension}")
            return None
